import os
import re
from typing import List, Dict
from docx import Document
from PyPDF2 import PdfReader


def _extract_table_text(table) -> str:
    """
    Convert DOCX table to readable text format.
    
    Args:
        table: python-docx Table object
        
    Returns:
        Formatted table text with headers and rows
    """
    rows = []
    for row in table.rows:
        cells = [cell.text.strip() for cell in row.cells]
        rows.append(' | '.join(cells))
    return '\n'.join(rows)


def parse_docx_structure(file_path: str) -> List[Dict]:
    """
    Parse DOCX file and extract heading hierarchy.
    
    Returns:
        List of sections with:
        - section_path: "5.2" (hierarchical numeric path)
        - heading_hierarchy: ["Policy", "Network Access"] (heading chain)
        - content: "Remote access to corporate systems..." (section text)
        - section_level: 2 (heading depth)
        - is_heading: True/False
        - is_table: True/False (NEW)
    """
    doc = Document(file_path)
    sections = []
    
    # Track heading hierarchy
    heading_stack = []  # [(level, title), ...]
    section_counter = [0, 0, 0, 0, 0]  # Track up to 5 levels
    has_first_heading = False  # Track if we've seen the first heading
    
    # Track current context for table insertion
    current_path = '0'
    current_hierarchy = ['Document Header']
    
    for para in doc.paragraphs:
        style = para.style.name if para.style else 'None'
        text = para.text.strip()
        
        if not text:
            continue
        
        if 'Heading' in style:
            # Extract heading level
            level_match = re.search(r'Heading (\d+)', style)
            level = int(level_match.group(1)) if level_match.group(1) else 1
            
            # Reset lower level counters
            for j in range(level, len(section_counter)):
                section_counter[j] = 0
            
            # Increment current level
            section_counter[level - 1] += 1
            
            # Build section path
            section_path = '.'.join(str(section_counter[j]) for j in range(level) if section_counter[j] > 0)
            
            # Update heading stack
            # Remove any headings at this level or deeper
            heading_stack = [(l, t) for l, t in heading_stack if l < level]
            heading_stack.append((level, text))
            
            # Build heading hierarchy
            heading_hierarchy = [t for l, t in heading_stack]
            
            # Update current context
            current_path = section_path
            current_hierarchy = heading_hierarchy
            
            has_first_heading = True
            
            sections.append({
                'section_path': section_path,
                'heading_hierarchy': heading_hierarchy,
                'content': text,
                'section_level': level,
                'is_heading': True,
                'is_table': False
            })
        else:
            # Content paragraph
            if has_first_heading:
                # We have a heading context
                current_path = '.'.join(str(section_counter[j]) for j in range(len(section_counter)) if section_counter[j] > 0)
                current_hierarchy = [t for l, t in heading_stack]
            else:
                # Before first heading - use "0" as section path
                current_path = '0'
                current_hierarchy = ['Document Header']
            
            sections.append({
                'section_path': current_path,
                'heading_hierarchy': current_hierarchy,
                'content': text,
                'section_level': len(heading_stack),
                'is_heading': False,
                'is_table': False
            })
    
    # NEW: Parse tables and insert at correct position
    for table in doc.tables:
        table_text = _extract_table_text(table)
        if table_text.strip():
            sections.append({
                'section_path': current_path,
                'heading_hierarchy': current_hierarchy,
                'content': table_text,
                'section_level': len(heading_stack),
                'is_heading': False,
                'is_table': True
            })
    
    return sections


def parse_pdf_structure(file_path: str) -> List[Dict]:
    """
    Parse PDF file and extract structure using font size heuristics.
    
    Returns:
        List of sections with:
        - section_path: "5.2" (hierarchical numeric path)
        - heading_hierarchy: ["Policy", "Network Access"] (heading chain)
        - content: "Remote access to corporate systems..." (section text)
        - section_level: 2 (heading depth)
        - is_heading: True/False
    """
    reader = PdfReader(file_path)
    sections = []
    
    # Track heading hierarchy
    heading_stack = []
    section_counter = [0, 0, 0, 0, 0]
    
    for page_num, page in enumerate(reader.pages):
        # Extract text with basic structure detection
        text = page.extract_text()
        if not text:
            continue
        
        lines = text.split('\n')
        
        for line in lines:
            line = line.strip()
            if not line:
                continue
            
            # Heuristic: Detect headings by characteristics
            is_heading = False
            level = 1
            
            # Check if line looks like a heading
            # 1. Short line (< 80 chars)
            # 2. Doesn't end with period
            # 3. May be bold (can't detect in plain text)
            # 4. May be centered (can't detect in plain text)
            if len(line) < 80 and not line.endswith('.'):
                # Additional heuristics
                # - All caps or title case
                # - Contains numbers at start (e.g., "1. Policy")
                # - Common heading words
                
                if re.match(r'^\d+\.?\s+[A-Z]', line):
                    # Numbered heading like "1. Policy" or "5.2 Network Access"
                    is_heading = True
                    # Extract level from number of dots
                    dots = line.count('.')
                    level = min(dots + 1, 5)
                elif line.isupper() or (line.istitle() and len(line.split()) <= 5):
                    # All caps or title case with few words
                    is_heading = True
                    level = 1
            
            if is_heading:
                # Reset lower level counters
                for j in range(level, len(section_counter)):
                    section_counter[j] = 0
                
                # Increment current level
                section_counter[level - 1] += 1
                
                # Build section path
                section_path = '.'.join(str(section_counter[j]) for j in range(level) if section_counter[j] > 0)
                
                # Update heading stack
                heading_stack = [(l, t) for l, t in heading_stack if l < level]
                heading_stack.append((level, line))
                
                # Build heading hierarchy
                heading_hierarchy = [t for l, t in heading_stack]
                
                sections.append({
                    'section_path': section_path,
                    'heading_hierarchy': heading_hierarchy,
                    'content': line,
                    'section_level': level,
                    'is_heading': True,
                    'page': page_num + 1
                })
            else:
                # Content paragraph
                current_path = '.'.join(str(section_counter[j]) for j in range(len(section_counter)) if section_counter[j] > 0)
                heading_hierarchy = [t for l, t in heading_stack]
                
                sections.append({
                    'section_path': current_path,
                    'heading_hierarchy': heading_hierarchy,
                    'content': line,
                    'section_level': len(heading_stack),
                    'is_heading': False,
                    'page': page_num + 1
                })
    
    return sections


def parse_document_structure(file_path: str) -> List[Dict]:
    """
    Parse document structure based on file extension.
    
    Args:
        file_path: Path to the document file
        
    Returns:
        List of sections with structure information
    """
    ext = os.path.splitext(file_path)[1].lower()
    
    if ext == '.docx':
        return parse_docx_structure(file_path)
    elif ext == '.pdf':
        return parse_pdf_structure(file_path)
    else:
        raise ValueError(f"Unsupported file format: {ext}")
